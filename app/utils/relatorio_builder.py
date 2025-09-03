from app import db
from sqlalchemy import text, select, and_, or_
import pandas as pd
from datetime import datetime
import os
import tempfile
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
import xlsxwriter
from docx import Document
from docx.shared import Inches, Pt
import json


class RelatorioBuilder:
    """Classe responsável por construir relatórios customizados"""

    def __init__(self):
        self.temp_dir = tempfile.mkdtemp()

    def processar_configuracao(self, config):
        """Processa a configuração e retorna os dados"""
        resultado = {
            'titulo': config.get('titulo', 'Relatório'),
            'secoes': []
        }

        for secao in config.get('secoes', []):
            dados_secao = self._processar_secao(secao)
            resultado['secoes'].append(dados_secao)

        return resultado

    def _processar_secao(self, secao):
        """Processa uma seção individual do relatório"""
        tipo = secao.get('tipo')

        if tipo == 'tabela':
            return self._processar_tabela(secao)
        elif tipo == 'grafico':
            return self._processar_grafico(secao)
        elif tipo == 'resumo':
            return self._processar_resumo(secao)
        elif tipo == 'texto':
            return {'tipo': 'texto', 'conteudo': secao.get('conteudo', '')}

    def _processar_tabela(self, config):
        """Busca dados e monta tabela"""
        fonte = config.get('fonte')
        campos = config.get('campos', [])
        filtros = config.get('filtros', [])
        ordenacao = config.get('ordenacao', {})
        limite = config.get('limite', 1000)

        # Montar query
        query = self._montar_query(fonte, campos, filtros, ordenacao, limite)

        # Executar query
        result = db.session.execute(text(query))
        registros = [dict(row) for row in result]

        # Formatar dados
        for registro in registros:
            for campo in campos:
                if campo['tipo'] == 'moeda':
                    valor = registro.get(campo['id'], 0)
                    registro[campo['id']] = f"R$ {valor:,.2f}" if valor else "R$ 0,00"
                elif campo['tipo'] == 'data':
                    valor = registro.get(campo['id'])
                    if valor:
                        registro[campo['id']] = valor.strftime('%d/%m/%Y')
                elif campo['tipo'] == 'percentual':
                    valor = registro.get(campo['id'], 0)
                    registro[campo['id']] = f"{valor:.2f}%"

        return {
            'tipo': 'tabela',
            'titulo': config.get('titulo', ''),
            'colunas': [c['nome'] for c in campos],
            'registros': registros,
            'total': len(registros)
        }

    def _processar_grafico(self, config):
        """Processa dados para gráfico"""
        fonte = config.get('fonte')
        tipo_grafico = config.get('tipo_grafico', 'barras')
        campo_x = config.get('campo_x')
        campo_y = config.get('campo_y')
        agregacao = config.get('agregacao', 'soma')

        # Query para dados do gráfico
        if agregacao == 'soma':
            func = 'SUM'
        elif agregacao == 'media':
            func = 'AVG'
        elif agregacao == 'contagem':
            func = 'COUNT'
        else:
            func = 'SUM'

        query = f"""
        SELECT 
            {campo_x} as categoria,
            {func}({campo_y}) as valor
        FROM BDG.{self._get_tabela_fonte(fonte)}
        WHERE DELETED_AT IS NULL
        GROUP BY {campo_x}
        ORDER BY 2 DESC
        LIMIT 10
        """

        result = db.session.execute(text(query))
        dados = [{'categoria': str(row[0]), 'valor': float(row[1] or 0)} for row in result]

        return {
            'tipo': 'grafico',
            'titulo': config.get('titulo', ''),
            'tipo_grafico': tipo_grafico,
            'dados': dados
        }

    def _processar_resumo(self, config):
        """Processa cards de resumo/KPIs"""
        metricas = []

        for metrica in config.get('metricas', []):
            valor = self._calcular_metrica(metrica)
            metricas.append({
                'titulo': metrica.get('titulo'),
                'valor': valor,
                'icone': metrica.get('icone', 'fa-chart-line'),
                'cor': metrica.get('cor', 'primary')
            })

        return {
            'tipo': 'resumo',
            'titulo': config.get('titulo', ''),
            'metricas': metricas
        }

    def _montar_query(self, fonte, campos, filtros, ordenacao, limite):
        """Monta query SQL dinamicamente"""
        tabela = self._get_tabela_fonte(fonte)

        # SELECT
        campos_select = ', '.join([c['id'] for c in campos])

        # WHERE
        condicoes = ['DELETED_AT IS NULL']
        for filtro in filtros:
            condicao = self._montar_condicao(filtro)
            if condicao:
                condicoes.append(condicao)

        where_clause = ' AND '.join(condicoes)

        # ORDER BY
        order_clause = ''
        if ordenacao:
            campo_ordem = ordenacao.get('campo', campos[0]['id'])
            direcao = ordenacao.get('direcao', 'ASC')
            order_clause = f"ORDER BY {campo_ordem} {direcao}"

        # Montar query completa
        query = f"""
        SELECT {campos_select}
        FROM BDG.{tabela}
        WHERE {where_clause}
        {order_clause}
        LIMIT {limite}
        """

        return query

    def _get_tabela_fonte(self, fonte):
        """Retorna o nome da tabela baseado na fonte"""
        mapa_tabelas = {
            'empresas': 'DCA_TB002_EMPRESAS_PARTICIPANTES',
            'distribuicao': 'DCA_TB005_DISTRIBUICAO',
            'metas': 'DCA_TB012_METAS_AVALIACAO',
            'contratos_distribuiveis': 'DCA_TB006_DISTRIBUIVEIS'
        }
        return mapa_tabelas.get(fonte, '')

    def _montar_condicao(self, filtro):
        """Monta condição WHERE baseada no filtro"""
        campo = filtro.get('campo')
        operador = filtro.get('operador')
        valor = filtro.get('valor')

        if not all([campo, operador, valor]):
            return None

        # Tratar diferentes operadores
        if operador == 'igual':
            return f"{campo} = '{valor}'"
        elif operador == 'diferente':
            return f"{campo} != '{valor}'"
        elif operador == 'maior':
            return f"{campo} > {valor}"
        elif operador == 'menor':
            return f"{campo} < {valor}"
        elif operador == 'contem':
            return f"{campo} LIKE '%{valor}%'"
        elif operador == 'entre':
            valores = valor.split(',')
            if len(valores) == 2:
                return f"{campo} BETWEEN '{valores[0]}' AND '{valores[1]}'"

        return None

    def _calcular_metrica(self, config):
        """Calcula valor de uma métrica"""
        query = config.get('query')
        if query:
            result = db.session.execute(text(query)).scalar()

            # Formatar valor baseado no tipo
            tipo = config.get('tipo_valor', 'numero')
            if tipo == 'moeda':
                return f"R$ {result:,.2f}" if result else "R$ 0,00"
            elif tipo == 'percentual':
                return f"{result:.1f}%" if result else "0%"
            else:
                return str(result or 0)

        return "0"

    def gerar(self, configuracao, formato, filtros, titulo):
        """Gera o arquivo do relatório no formato especificado"""
        # Processar dados
        dados = self.processar_configuracao(configuracao)

        # Aplicar filtros globais
        if filtros:
            dados['filtros_aplicados'] = filtros

        # Gerar arquivo baseado no formato
        if formato == 'pdf':
            return self._gerar_pdf(dados, titulo)
        elif formato == 'excel':
            return self._gerar_excel(dados, titulo)
        elif formato == 'word':
            return self._gerar_word(dados, titulo)

    def _gerar_pdf(self, dados, titulo):
        """Gera relatório em PDF"""
        nome_arquivo = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        caminho = os.path.join(self.temp_dir, nome_arquivo)

        # Criar documento
        doc = SimpleDocTemplate(caminho, pagesize=A4)
        story = []
        styles = getSampleStyleSheet()

        # Título
        titulo_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#2c3e50')
        )
        story.append(Paragraph(titulo, titulo_style))
        story.append(Spacer(1, 0.2 * inch))

        # Data de geração
        data_style = ParagraphStyle(
            'DataStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey
        )
        story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", data_style))
        story.append(Spacer(1, 0.3 * inch))

        # Processar seções
        for secao in dados.get('secoes', []):
            if secao['tipo'] == 'tabela':
                self._adicionar_tabela_pdf(story, secao, styles)
            elif secao['tipo'] == 'texto':
                story.append(Paragraph(secao['conteudo'], styles['Normal']))
                story.append(Spacer(1, 0.2 * inch))
            elif secao['tipo'] == 'resumo':
                self._adicionar_resumo_pdf(story, secao, styles)

        # Gerar PDF
        doc.build(story)

        return {
            'nome': nome_arquivo,
            'caminho': caminho,
            'mimetype': 'application/pdf'
        }

    def _adicionar_tabela_pdf(self, story, secao, styles):
        """Adiciona tabela ao PDF"""
        if secao.get('titulo'):
            story.append(Paragraph(secao['titulo'], styles['Heading2']))
            story.append(Spacer(1, 0.1 * inch))

        # Preparar dados da tabela
        data = [secao['colunas']]  # Cabeçalho

        for registro in secao['registros'][:50]:  # Limitar a 50 registros no PDF
            linha = []
            for coluna in secao['colunas']:
                # Buscar valor correspondente no registro
                valor = ''
                for campo, val in registro.items():
                    if campo in str(coluna).lower():
                        valor = str(val)
                        break
                linha.append(valor)
            data.append(linha)

        # Criar tabela
        t = Table(data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(t)
        story.append(Spacer(1, 0.3 * inch))

        # Adicionar nota se há mais registros
        if len(secao['registros']) > 50:
            nota = Paragraph(
                f"* Mostrando apenas os primeiros 50 de {len(secao['registros'])} registros",
                styles['Italic']
            )
            story.append(nota)
            story.append(Spacer(1, 0.3 * inch))

    def _adicionar_resumo_pdf(self, story, secao, styles):
        """Adiciona cards de resumo ao PDF"""
        if secao.get('titulo'):
            story.append(Paragraph(secao['titulo'], styles['Heading2']))
            story.append(Spacer(1, 0.1 * inch))

        # Criar tabela para os KPIs
        data = []
        linha = []

        for i, metrica in enumerate(secao['metricas']):
            kpi_text = f"<b>{metrica['titulo']}</b><br/>{metrica['valor']}"
            linha.append(Paragraph(kpi_text, styles['Normal']))

            # Quebrar linha a cada 3 KPIs
            if (i + 1) % 3 == 0:
                data.append(linha)
                linha = []

        if linha:  # Adicionar última linha incompleta
            data.append(linha)

        if data:
            t = Table(data, colWidths=[2 * inch] * 3)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#ecf0f1')),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('GRID', (0, 0), (-1, -1), 1, colors.white),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 20),
                ('TOPPADDING', (0, 0), (-1, -1), 20)
            ]))
            story.append(t)
            story.append(Spacer(1, 0.3 * inch))

    def _gerar_excel(self, dados, titulo):
        """Gera relatório em Excel"""
        nome_arquivo = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        caminho = os.path.join(self.temp_dir, nome_arquivo)

        # Criar workbook
        workbook = xlsxwriter.Workbook(caminho)

        # Formatos
        titulo_format = workbook.add_format({
            'bold': True,
            'font_size': 16,
            'font_color': '#2c3e50'
        })

        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#34495e',
            'font_color': 'white',
            'border': 1
        })

        cell_format = workbook.add_format({
            'border': 1
        })

        money_format = workbook.add_format({
            'border': 1,
            'num_format': 'R$ #,##0.00'
        })

        # Criar abas para cada seção tipo tabela
        for i, secao in enumerate(dados.get('secoes', [])):
            if secao['tipo'] == 'tabela':
                # Nome da aba (máximo 31 caracteres)
                nome_aba = (secao.get('titulo', f'Tabela_{i + 1}'))[:31]
                worksheet = workbook.add_worksheet(nome_aba)

                # Título
                worksheet.write(0, 0, titulo, titulo_format)
                worksheet.write(1, 0, f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

                # Cabeçalhos
                row = 3
                for col, header in enumerate(secao['colunas']):
                    worksheet.write(row, col, header, header_format)

                # Dados
                row = 4
                for registro in secao['registros']:
                    col = 0
                    for campo in secao['colunas']:
                        valor = ''
                        # Buscar valor no registro
                        for k, v in registro.items():
                            if campo.lower() in k.lower():
                                valor = v
                                break

                        # Aplicar formato apropriado
                        if 'R$' in str(valor):
                            # Remover formatação e converter para número
                            valor_num = float(str(valor).replace('R$', '').replace('.', '').replace(',', '.'))
                            worksheet.write(row, col, valor_num, money_format)
                        else:
                            worksheet.write(row, col, valor, cell_format)
                        col += 1
                    row += 1

                # Ajustar largura das colunas
                for col in range(len(secao['colunas'])):
                    worksheet.set_column(col, col, 20)

        # Adicionar aba de resumo se houver
        resumos = [s for s in dados.get('secoes', []) if s['tipo'] == 'resumo']
        if resumos:
            worksheet = workbook.add_worksheet('Resumo')
            worksheet.write(0, 0, 'Resumo Executivo', titulo_format)

            row = 2
            for resumo in resumos:
                if resumo.get('titulo'):
                    worksheet.write(row, 0, resumo['titulo'], header_format)
                    row += 1

                for metrica in resumo['metricas']:
                    worksheet.write(row, 0, metrica['titulo'])
                    worksheet.write(row, 1, metrica['valor'])
                    row += 1
                row += 1

        workbook.close()

        return {
            'nome': nome_arquivo,
            'caminho': caminho,
            'mimetype': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        }

    def _gerar_word(self, dados, titulo):
        """Gera relatório em Word"""
        nome_arquivo = f"relatorio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        caminho = os.path.join(self.temp_dir, nome_arquivo)

        # Criar documento
        doc = Document()

        # Título
        doc.add_heading(titulo, 0)
        doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph()

        # Processar seções
        for secao in dados.get('secoes', []):
            if secao['tipo'] == 'tabela':
                if secao.get('titulo'):
                    doc.add_heading(secao['titulo'], level=1)

                # Criar tabela
                table = doc.add_table(rows=1, cols=len(secao['colunas']))
                table.style = 'Light Grid Accent 1'

                # Cabeçalhos
                hdr_cells = table.rows[0].cells
                for i, coluna in enumerate(secao['colunas']):
                    hdr_cells[i].text = coluna

                # Dados (limitar a 100 registros)
                for registro in secao['registros'][:100]:
                    row_cells = table.add_row().cells
                    for i, coluna in enumerate(secao['colunas']):
                        # Buscar valor
                        valor = ''
                        for k, v in registro.items():
                            if coluna.lower() in k.lower():
                                valor = str(v)
                                break
                        row_cells[i].text = valor

                if len(secao['registros']) > 100:
                    doc.add_paragraph(
                        f"* Mostrando apenas os primeiros 100 de {len(secao['registros'])} registros"
                    )

                doc.add_paragraph()

            elif secao['tipo'] == 'texto':
                doc.add_paragraph(secao['conteudo'])

            elif secao['tipo'] == 'resumo':
                if secao.get('titulo'):
                    doc.add_heading(secao['titulo'], level=1)

                for metrica in secao['metricas']:
                    p = doc.add_paragraph()
                    p.add_run(f"{metrica['titulo']}: ").bold = True
                    p.add_run(metrica['valor'])

                doc.add_paragraph()

        # Salvar
        doc.save(caminho)

        return {
            'nome': nome_arquivo,
            'caminho': caminho,
            'mimetype': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        }