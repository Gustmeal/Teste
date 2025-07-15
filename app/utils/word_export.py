from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import tempfile


def export_to_word(dados, colunas, titulo):
    """
    Exporta dados para um arquivo Word

    Args:
        dados: Lista de dicionários com os dados a serem exportados
        colunas: Lista de nomes das colunas na ordem desejada
        titulo: Título do relatório

    Returns:
        BytesIO object com o arquivo Word
    """
    # Criar um novo documento Word
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Adicionar título
    title = doc.add_heading(titulo, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adicionar data de geração
    data_geracao = doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")
    data_geracao.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Adicionar linha em branco
    doc.add_paragraph()

    # Verificar se há dados
    if not dados:
        doc.add_paragraph("Nenhum dado disponível para exportação.")
        # Salvar em BytesIO
        import io
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    # Criar tabela
    table = doc.add_table(rows=1, cols=len(colunas))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Adicionar cabeçalho
    hdr_cells = table.rows[0].cells
    for i, coluna in enumerate(colunas):
        hdr_cells[i].text = coluna
        # Formatar cabeçalho
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Cor de fundo do cabeçalho
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        tcVAlign = OxmlElement('w:shd')
        tcVAlign.set(qn('w:val'), 'clear')
        tcVAlign.set(qn('w:color'), 'auto')
        tcVAlign.set(qn('w:fill'), '5B52E5')  # Cor roxo primário
        tcPr.append(tcVAlign)

    # Adicionar dados
    for row_data in dados:
        row_cells = table.add_row().cells
        for i, coluna in enumerate(colunas):
            value = row_data.get(coluna, '')
            # Converter valor para string
            if value is None:
                value = ''
            elif isinstance(value, (int, float)):
                if isinstance(value, float):
                    value = f"{value:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
                else:
                    value = str(value)
            else:
                value = str(value)

            row_cells[i].text = value
            # Alinhar números à direita
            if any(char.isdigit() for char in value) and (
                    value.startswith('R$') or value.endswith('%') or value.replace('.', '').replace(',', '').isdigit()):
                for paragraph in row_cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Ajustar largura das colunas
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            # Definir largura mínima
            cell.width = Inches(1.5)

    # Adicionar rodapé
    doc.add_paragraph()
    footer = doc.add_paragraph("Sistema GEINC - Gestão Integrada de Credenciamento")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(10)
    footer.runs[0].font.italic = True

    # Salvar em BytesIO
    import io
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return output