from flask import Blueprint, render_template, request, flash, send_file, redirect, url_for
from flask_login import login_required
from datetime import datetime
import os
import tempfile
import uuid
import mammoth
import pandas as pd
from openpyxl import Workbook
from PyPDF2 import PdfReader
import docx
import xlsxwriter
import io

converter_bp = Blueprint('converter', __name__, url_prefix='/converter')


@converter_bp.context_processor
def inject_current_year():
    return {'current_year': datetime.utcnow().year}


@converter_bp.route('/')
@login_required
def index():
    return render_template('converter/index.html')


@converter_bp.route('/convert', methods=['POST'])
@login_required
def convert():
    if 'file' not in request.files:
        flash('Nenhum arquivo enviado.', 'danger')
        return redirect(url_for('converter.index'))

    file = request.files['file']

    if file.filename == '':
        flash('Nenhum arquivo selecionado.', 'danger')
        return redirect(url_for('converter.index'))

    # Obter o formato de conversão selecionado
    conversion_type = request.form.get('conversion_type')
    if not conversion_type:
        flash('Tipo de conversão não selecionado.', 'danger')
        return redirect(url_for('converter.index'))

    # Determinar extensões
    source_ext = os.path.splitext(file.filename)[1].lower()

    # Criar diretório temporário para trabalhar com os arquivos
    temp_dir = tempfile.gettempdir()
    temp_filename = str(uuid.uuid4())
    temp_source_path = os.path.join(temp_dir, f"{temp_filename}{source_ext}")

    # Salvar arquivo enviado
    file.save(temp_source_path)

    try:
        # Processar conversão baseado no tipo
        if conversion_type == 'word_to_pdf':
            if source_ext not in ['.docx', '.doc']:
                flash('Por favor, carregue um arquivo Word (.docx ou .doc) para esta conversão.', 'danger')
                return redirect(url_for('converter.index'))

            # Implementação da conversão Word para PDF
            result_file, download_name = convert_word_to_pdf(temp_source_path, file.filename)

        elif conversion_type == 'word_to_excel':
            if source_ext not in ['.docx', '.doc']:
                flash('Por favor, carregue um arquivo Word (.docx ou .doc) para esta conversão.', 'danger')
                return redirect(url_for('converter.index'))

            # Implementação da conversão Word para Excel
            result_file, download_name = convert_word_to_excel(temp_source_path, file.filename)

        elif conversion_type == 'pdf_to_excel':
            if source_ext != '.pdf':
                flash('Por favor, carregue um arquivo PDF para esta conversão.', 'danger')
                return redirect(url_for('converter.index'))

            # Implementação da conversão PDF para Excel
            result_file, download_name = convert_pdf_to_excel(temp_source_path, file.filename)

        else:
            flash('Tipo de conversão não suportado.', 'danger')
            return redirect(url_for('converter.index'))

        # Retornar o arquivo convertido
        return send_file(
            result_file,
            as_attachment=True,
            download_name=download_name,
            max_age=0
        )

    except Exception as e:
        flash(f'Erro durante a conversão: {str(e)}', 'danger')
        return redirect(url_for('converter.index'))
    finally:
        # Limpar arquivos temporários
        try:
            if os.path.exists(temp_source_path):
                os.remove(temp_source_path)
        except:
            pass


def convert_word_to_pdf(word_path, original_filename):
    """Converte arquivo Word para PDF"""
    from docx2pdf import convert

    base_name = os.path.splitext(original_filename)[0]
    temp_dir = tempfile.gettempdir()
    output_path = os.path.join(temp_dir, f"{base_name}.pdf")

    # Converter o documento
    convert(word_path, output_path)

    return output_path, f"{base_name}.pdf"


def convert_word_to_excel(word_path, original_filename):
    """Extrai tabelas de Word para Excel"""
    doc = docx.Document(word_path)
    base_name = os.path.splitext(original_filename)[0]

    # Criar workbook Excel
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet("Texto Extraído")

    # Extrair texto e tabelas
    row_num = 0
    for para in doc.paragraphs:
        worksheet.write(row_num, 0, para.text)
        row_num += 1

    # Se houver tabelas, adicionar cada uma em uma nova planilha
    for i, table in enumerate(doc.tables):
        sheet_name = f"Tabela {i + 1}"
        table_sheet = workbook.add_worksheet(sheet_name)

        # Extrair dados da tabela
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                table_sheet.write(r_idx, c_idx, cell.text)

    workbook.close()
    output.seek(0)

    # Retornar bytes do Excel
    return output, f"{base_name}.xlsx"


def convert_pdf_to_excel(pdf_path, original_filename):
    """Extrai texto e tabelas de PDF para Excel"""
    # Abrir o PDF
    pdf = PdfReader(pdf_path)
    base_name = os.path.splitext(original_filename)[0]

    # Criar workbook Excel
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet("Conteúdo PDF")

    # Extrair texto de cada página
    row_num = 0
    for page_num in range(len(pdf.pages)):
        page = pdf.pages[page_num]
        text = page.extract_text()

        # Adicionar número da página
        worksheet.write(row_num, 0, f"Página {page_num + 1}")
        row_num += 1

        # Adicionar texto da página
        lines = text.split('\n')
        for line in lines:
            worksheet.write(row_num, 0, line)
            row_num += 1

        # Espaço entre páginas
        row_num += 1

    workbook.close()
    output.seek(0)

    # Retornar bytes do Excel
    return output, f"{base_name}.xlsx"